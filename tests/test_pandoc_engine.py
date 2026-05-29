"""Tests for PandocEngine conversions."""

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from omnicon.core.job import ConversionJob
from omnicon.engines.pandoc_engine import PandocEngine


@pytest.fixture
def engine() -> PandocEngine:
    return PandocEngine()


@pytest.fixture
def sample_md(tmp_path: Path) -> Path:
    """Create a minimal Markdown file for testing."""
    path = tmp_path / "sample.md"
    path.write_text(
        "# Hello OmniCon\n\nThis is a **test** paragraph.\n\n- Item one\n- Item two\n",
        encoding="utf-8",
    )
    return path


@pytest.fixture
def sample_html(tmp_path: Path) -> Path:
    """Create a minimal HTML file for testing."""
    path = tmp_path / "sample.html"
    path.write_text(
        "<!DOCTYPE html><html><body><h1>Heading</h1><p>Paragraph text.</p></body></html>",
        encoding="utf-8",
    )
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX file using python-docx."""
    from docx import Document

    doc = Document()
    doc.add_heading("DOCX Heading", level=1)
    doc.add_paragraph("DOCX paragraph content.")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


def make_job(input_path: Path, output_format: str, output_dir: Path) -> ConversionJob:
    return ConversionJob(
        input_path=input_path,
        output_format=output_format,
        output_dir=output_dir,
    )


def _mock_convert_file(
    source_file: str,
    to: str,
    format: str | None = None,
    outputfile: str | None = None,
    extra_args: list[str] | None = None,
) -> str:
    """Fake pypandoc.convert_file that creates a plausible output file."""
    if outputfile:
        out = Path(outputfile)
        if to == "docx":
            # Write a minimal DOCX-like file (just needs to be non-empty)
            from docx import Document

            doc = Document()
            doc.add_paragraph("Converted content")
            doc.save(outputfile)
        elif to == "html":
            out.write_text(
                "<html><body><h1>Converted</h1></body></html>",
                encoding="utf-8",
            )
        elif to == "plain":
            out.write_text("Converted plain text content", encoding="utf-8")
        elif to == "markdown":
            out.write_text("# Converted\n\nContent", encoding="utf-8")
        elif to == "pdf":
            out.write_bytes(b"%PDF-1.4 fake")
        else:
            out.write_text("converted", encoding="utf-8")
    return ""


# --- can_convert tests ---


def test_can_convert_supported(engine: PandocEngine) -> None:
    assert engine.can_convert("md", "pdf") is True
    assert engine.can_convert("md", "docx") is True
    assert engine.can_convert("md", "html") is True
    assert engine.can_convert("docx", "md") is True
    assert engine.can_convert("html", "md") is True
    assert engine.can_convert("md", "txt") is True


def test_can_convert_unsupported(engine: PandocEngine) -> None:
    assert engine.can_convert("pdf", "md") is False
    assert engine.can_convert("docx", "pdf") is False
    assert engine.can_convert("html", "pdf") is False


def test_can_convert_case_insensitive(engine: PandocEngine) -> None:
    assert engine.can_convert("MD", "PDF") is True
    assert engine.can_convert("Md", "Docx") is True


def test_priority(engine: PandocEngine) -> None:
    assert engine.priority == 30


# --- md -> docx tests ---


def test_md_to_docx(engine: PandocEngine, sample_md: Path, tmp_path: Path) -> None:
    """Markdown to DOCX conversion produces a valid .docx file."""
    job = make_job(sample_md, "docx", tmp_path)

    mock_pypandoc = MagicMock()
    mock_pypandoc.convert_file = _mock_convert_file

    with patch.dict("sys.modules", {"pypandoc": mock_pypandoc}):
        result = engine.convert(job)

    assert result.exists()
    assert result.suffix == ".docx"
    assert result.stat().st_size > 0
    assert job.progress == 100


# --- md -> html tests ---


def test_md_to_html(engine: PandocEngine, sample_md: Path, tmp_path: Path) -> None:
    """Markdown to HTML conversion produces a valid .html file."""
    job = make_job(sample_md, "html", tmp_path)

    mock_pypandoc = MagicMock()
    mock_pypandoc.convert_file = _mock_convert_file

    with patch.dict("sys.modules", {"pypandoc": mock_pypandoc}):
        result = engine.convert(job)

    assert result.exists()
    assert result.suffix == ".html"
    content = result.read_text(encoding="utf-8")
    assert "<html>" in content.lower()
    assert job.progress == 100


# --- md -> txt tests ---


def test_md_to_txt(engine: PandocEngine, sample_md: Path, tmp_path: Path) -> None:
    """Markdown to plain text conversion produces a valid .txt file."""
    job = make_job(sample_md, "txt", tmp_path)

    mock_pypandoc = MagicMock()
    mock_pypandoc.convert_file = _mock_convert_file

    with patch.dict("sys.modules", {"pypandoc": mock_pypandoc}):
        result = engine.convert(job)

    assert result.exists()
    assert result.suffix == ".txt"
    content = result.read_text(encoding="utf-8")
    assert len(content) > 0
    assert job.progress == 100


# --- docx -> md tests ---


def test_docx_to_md(engine: PandocEngine, sample_docx: Path, tmp_path: Path) -> None:
    """DOCX to Markdown conversion produces a valid .md file."""
    job = make_job(sample_docx, "md", tmp_path)

    mock_pypandoc = MagicMock()
    mock_pypandoc.convert_file = _mock_convert_file

    with patch.dict("sys.modules", {"pypandoc": mock_pypandoc}):
        result = engine.convert(job)

    assert result.exists()
    assert result.suffix == ".md"
    content = result.read_text(encoding="utf-8")
    assert len(content) > 0
    assert job.progress == 100


# --- html -> md tests ---


def test_html_to_md(engine: PandocEngine, sample_html: Path, tmp_path: Path) -> None:
    """HTML to Markdown conversion produces a valid .md file."""
    job = make_job(sample_html, "md", tmp_path)

    mock_pypandoc = MagicMock()
    mock_pypandoc.convert_file = _mock_convert_file

    with patch.dict("sys.modules", {"pypandoc": mock_pypandoc}):
        result = engine.convert(job)

    assert result.exists()
    assert result.suffix == ".md"
    content = result.read_text(encoding="utf-8")
    assert len(content) > 0
    assert job.progress == 100


# --- md -> pdf tests ---


def test_md_to_pdf_uses_weasyprint_engine(
    engine: PandocEngine, sample_md: Path, tmp_path: Path
) -> None:
    """Markdown to PDF passes --pdf-engine=weasyprint to pypandoc."""
    job = make_job(sample_md, "pdf", tmp_path)

    captured_args: dict[str, list[str]] = {}

    def capturing_convert_file(
        source_file: str,
        to: str,
        format: str | None = None,
        outputfile: str | None = None,
        extra_args: list[str] | None = None,
    ) -> str:
        captured_args["extra_args"] = extra_args or []
        # Create the output file
        if outputfile:
            Path(outputfile).write_bytes(b"%PDF-1.4 fake")
        return ""

    mock_pypandoc = MagicMock()
    mock_pypandoc.convert_file = capturing_convert_file

    with patch.dict("sys.modules", {"pypandoc": mock_pypandoc}):
        result = engine.convert(job)

    assert result.exists()
    assert "--pdf-engine=weasyprint" in captured_args["extra_args"]
    assert job.progress == 100


# --- error handling tests ---


def test_unsupported_route_raises(engine: PandocEngine, tmp_path: Path) -> None:
    """Unsupported route raises UnsupportedRouteError."""
    from omnicon.engines.base import UnsupportedRouteError

    fake_input = tmp_path / "fake.pdf"
    fake_input.write_bytes(b"fake")
    job = make_job(fake_input, "docx", tmp_path)

    with pytest.raises(UnsupportedRouteError):
        engine.convert(job)
