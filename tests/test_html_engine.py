"""Tests for HTMLEngine conversions."""

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from omnicon.core.job import ConversionJob
from omnicon.engines.html_engine import HTMLEngine


@pytest.fixture
def engine() -> HTMLEngine:
    return HTMLEngine()


@pytest.fixture
def sample_html(tmp_path: Path) -> Path:
    """Create a minimal HTML file for testing."""
    path = tmp_path / "sample.html"
    path.write_text(
        "<!DOCTYPE html><html><head><meta charset='utf-8'></head>"
        "<body><h1>Hello OmniCon</h1><p>Test paragraph.</p></body></html>",
        encoding="utf-8",
    )
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX file using python-docx."""
    from docx import Document

    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("This is a test paragraph for OmniCon.")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


def make_job(input_path: Path, output_format: str, output_dir: Path) -> ConversionJob:
    return ConversionJob(
        input_path=input_path,
        output_format=output_format,
        output_dir=output_dir,
    )


# --- can_convert tests ---


def test_can_convert_supported(engine: HTMLEngine) -> None:
    assert engine.can_convert("html", "pdf") is True
    assert engine.can_convert("docx", "html") is True
    assert engine.can_convert("html", "txt") is True


def test_can_convert_unsupported(engine: HTMLEngine) -> None:
    assert engine.can_convert("pdf", "html") is False
    assert engine.can_convert("html", "docx") is False
    assert engine.can_convert("txt", "html") is False


def test_can_convert_case_insensitive(engine: HTMLEngine) -> None:
    assert engine.can_convert("HTML", "PDF") is True
    assert engine.can_convert("Docx", "Html") is True


def test_priority(engine: HTMLEngine) -> None:
    assert engine.priority == 20


# --- html -> txt tests ---


def test_html_to_text(engine: HTMLEngine, sample_html: Path, tmp_path: Path) -> None:
    job = make_job(sample_html, "txt", tmp_path)
    result = engine.convert(job)

    assert result.exists()
    assert result.suffix == ".txt"
    content = result.read_text(encoding="utf-8")
    assert "Hello OmniCon" in content
    assert "Test paragraph" in content
    # Tags should be stripped
    assert "<h1>" not in content
    assert "<p>" not in content
    assert job.progress == 100


def test_html_to_text_whitespace_normalization(engine: HTMLEngine, tmp_path: Path) -> None:
    """Verify that excessive whitespace is collapsed to single spaces."""
    html_path = tmp_path / "spaced.html"
    html_path.write_text(
        "<html><body>  Word1   Word2\n\n\nWord3  </body></html>",
        encoding="utf-8",
    )
    job = make_job(html_path, "txt", tmp_path)
    result = engine.convert(job)

    content = result.read_text(encoding="utf-8")
    assert "Word1 Word2 Word3" in content


# --- docx -> html tests ---


def test_docx_to_html(engine: HTMLEngine, sample_docx: Path, tmp_path: Path) -> None:
    job = make_job(sample_docx, "html", tmp_path)
    result = engine.convert(job)

    assert result.exists()
    assert result.suffix == ".html"
    content = result.read_text(encoding="utf-8")
    assert "<!DOCTYPE html>" in content
    assert "Test Heading" in content
    assert "test paragraph" in content.lower()
    assert job.progress == 100


# --- html -> pdf tests (mocked) ---


def test_html_to_pdf_mocked(engine: HTMLEngine, sample_html: Path, tmp_path: Path) -> None:
    """Test HTML to PDF route with mocked WeasyPrint."""
    job = make_job(sample_html, "pdf", tmp_path)
    output_path = job.expected_output_path

    mock_html_cls = MagicMock()
    mock_html_instance = MagicMock()
    mock_html_cls.return_value = mock_html_instance

    # Simulate write_pdf creating the output file
    def fake_write_pdf(path: str) -> None:
        Path(path).write_bytes(b"%PDF-1.4 fake")

    mock_html_instance.write_pdf.side_effect = fake_write_pdf

    with patch.dict("sys.modules", {"weasyprint": MagicMock(HTML=mock_html_cls)}):
        # Re-import to pick up the mock
        from omnicon.engines.html_engine import HTMLEngine as FreshEngine

        fresh_engine = FreshEngine()
        result = fresh_engine.convert(job)

    assert result == output_path
    assert job.progress == 100


# --- error handling tests ---


def test_unsupported_route_raises(engine: HTMLEngine, tmp_path: Path) -> None:
    """Unsupported route raises UnsupportedRouteError."""
    from omnicon.engines.base import UnsupportedRouteError

    fake_input = tmp_path / "fake.pdf"
    fake_input.write_bytes(b"fake")
    job = make_job(fake_input, "docx", tmp_path)

    with pytest.raises(UnsupportedRouteError):
        engine.convert(job)
